"""
附件解析测试：真实生成 PDF/DOCX/XLSX 字节流验证解析与状态收敛，全部离线。

Run: .venv/bin/python -m pytest tests/test_attachment_parsers.py -v
"""

import datetime
import io
import shutil

import fitz  # pymupdf
import openpyxl
import pytest
from docx import Document

from sufe_qa.ingest.attachment_parsers import _detect_format, parse_attachment

OLE2_HEAD = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512

_LONG_PAGE = (
    "上海财经大学研究生学籍管理办法。第一条 为规范研究生学籍管理，"
    "根据教育部有关规定，结合学校实际，制定本办法。第二条 本办法适用于全校研究生。"
    "第三条 研究生学籍管理工作坚持公平公正原则，各单位应当严格遵照执行，"
    "并将执行情况按学期报送研究生院备案。"
)


def _make_pdf(page_texts: list[str]) -> bytes:
    """生成真实 PDF 字节流；空串页面不写字（模拟扫描件页面）。中文需 china-s 字体。

    用 insert_textbox 而非 insert_text：后者不换行，长行超出页宽会被裁掉。
    """
    doc = fitz.open()
    for t in page_texts:
        page = doc.new_page()
        if t:
            page.insert_textbox(fitz.Rect(72, 72, 520, 760), t, fontname="china-s")
    data = doc.tobytes()
    doc.close()
    return data


def _make_docx() -> bytes:
    d = Document()
    d.add_heading("上海财经大学研究生管理办法", level=1)
    d.add_paragraph("第一条 为规范研究生管理，制定本办法。")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "标准"
    table.cell(1, 0).text = "学制"
    table.cell(1, 1).text = "两年"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_xlsx() -> bytes:
    """两个工作表：含合并单元格、全空行、连续重复表头、日期单元格。"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "资助标准"
    ws.append(["类别", "金额", "备注"])
    ws.append(["奖学金", 5000, "按年发放"])
    ws.append([None, None, None])  # 全空行，应跳过
    ws.append(["助学金", 3000, "按月发放"])
    ws["A5"] = "合计说明"
    ws.merge_cells("A5:B5")  # 合并区域，B5 应填充左上角值
    ws2 = wb.create_sheet("申请流程")
    ws2.append(["步骤", "说明"])
    ws2.append(["步骤", "说明"])  # 连续重复表头，应去重
    ws2.append(["1", "网上申报"])
    ws2.append(["发布日期", datetime.date(2025, 3, 1)])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------- 格式识别 ----------


def test_detect_format_by_extension_and_magic():
    assert _detect_format("a.pdf", b"%PDF-1.7 rest") == "pdf"
    assert _detect_format("a.docx", _make_docx()) == "docx"  # zip 魔数 + .docx 一致
    assert _detect_format("a.doc", OLE2_HEAD) == "doc"
    assert _detect_format("a.xls", OLE2_HEAD) == "xls"
    assert _detect_format("a.pptx", b"PK\x03\x04 junk") == "pptx"
    assert _detect_format("a.txt", b"hello") == "unknown"


def test_detect_format_magic_wins_over_extension():
    # docx 字节（zip 魔数）却叫 .pdf：以魔数为准，窥探包内结构识别为 docx
    assert _detect_format("weird.pdf", _make_docx()) == "docx"
    # OLE2 魔数却叫 .docx：以魔数为准归入 legacy，但扩展名无法细分
    assert _detect_format("weird.docx", OLE2_HEAD) == "unknown"


# ---------- PDF ----------


def test_parse_pdf_ok():
    r = parse_attachment("办法.pdf", _make_pdf([_LONG_PAGE, _LONG_PAGE]))
    assert r.fmt == "pdf"
    assert r.parse_status == "ok"
    assert r.page_count == 2
    assert "第一条" in r.text
    assert r.char_count == len(r.text) > 0


def test_parse_pdf_scanned_all_empty_pages():
    r = parse_attachment("扫描件.pdf", _make_pdf(["", "", ""]))
    assert r.fmt == "pdf"
    assert r.parse_status == "scanned_pdf"
    assert r.page_count == 3
    assert r.text == ""
    assert any("疑似扫描件" in n for n in r.notes)


def test_parse_pdf_partial_scanned_notes_first_sparse_page():
    # 第 1 页有文本、后两页空白：平均 < 50 字符判扫描件，备注定位首个稀疏页
    r = parse_attachment("半扫描.pdf", _make_pdf([_LONG_PAGE, "", ""]))
    assert r.parse_status == "scanned_pdf"
    assert r.page_count == 3
    assert "第一条" in r.text  # 已提取文本仍返回
    assert any("第2页起疑似扫描件" in n for n in r.notes)


def test_parse_pdf_corrupt_returns_parse_failed_without_raising():
    r = parse_attachment("坏文件.pdf", b"%PDF-1.4 garbage-truncated")
    assert r.fmt == "pdf"
    assert r.parse_status == "parse_failed"
    assert r.text == ""
    assert r.notes  # 记录原因


def test_parse_pdf_over_page_limit_is_explicitly_quarantined_from_index():
    r = parse_attachment("超大档案.pdf", _make_pdf([_LONG_PAGE] * 301))
    assert r.fmt == "pdf"
    assert r.page_count == 301
    assert r.parse_status == "pdf_too_large"
    assert r.text == ""
    assert any("安全解析上限" in note for note in r.notes)


# ---------- DOCX ----------


def test_parse_docx_ok_with_heading_and_table():
    r = parse_attachment("办法.docx", _make_docx())
    assert r.fmt == "docx"
    assert r.parse_status == "ok"
    assert "## 上海财经大学研究生管理办法" in r.text  # 标题 markdown 化
    assert "第一条 为规范研究生管理" in r.text
    assert "项目 | 标准" in r.text  # 表格沿用 " | " 拼接
    assert "学制 | 两年" in r.text
    assert r.char_count == len(r.text)


def test_parse_docx_corrupt_returns_parse_failed():
    # zip 魔数但内容不是合法 docx
    r = parse_attachment("坏.docx", b"PK\x03\x04 not-a-real-zip")
    assert r.fmt == "docx"
    assert r.parse_status == "parse_failed"
    assert r.text == ""


# ---------- DOC（legacy） ----------


def test_parse_doc_legacy_without_libreoffice(monkeypatch):
    # 无 LibreOffice/soffice 时放弃解析，不抛异常
    monkeypatch.setattr(shutil, "which", lambda _cmd: None)
    r = parse_attachment("旧文件.doc", OLE2_HEAD)
    assert r.fmt == "doc"
    assert r.parse_status == "legacy_doc_unparsed"
    assert r.text == ""
    assert r.char_count == 0
    assert any("LibreOffice" in n for n in r.notes)


def test_parse_doc_conversion_failure(monkeypatch, tmp_path):
    # LibreOffice 存在但转换失败（非零退出且无产出）→ legacy_doc_unparsed
    monkeypatch.setattr(
        shutil, "which", lambda cmd: "/usr/bin/false" if cmd == "libreoffice" else None
    )
    r = parse_attachment("旧文件.doc", OLE2_HEAD)
    assert r.fmt == "doc"
    assert r.parse_status == "legacy_doc_unparsed"
    assert r.text == ""


# ---------- XLSX ----------


def test_parse_xlsx_ok_two_sheets_merged_and_dedup():
    r = parse_attachment("资助.xlsx", _make_xlsx())
    assert r.fmt == "xlsx"
    assert r.parse_status == "ok"
    assert r.sheet_count == 2
    assert "## 资助标准" in r.text
    assert "## 申请流程" in r.text
    # 全空行已跳过：奖学金行与助学金行相邻
    assert "奖学金 | 5000 | 按年发放\n助学金 | 3000 | 按月发放" in r.text
    # 合并单元格 B5 填充左上角值
    assert "合计说明 | 合计说明" in r.text
    # 连续重复表头去重
    assert r.text.count("步骤 | 说明") == 1
    # 日期显示为 ISO
    assert "发布日期 | 2025-03-01" in r.text
    assert r.char_count == len(r.text)


def test_parse_xlsx_truncates_huge_sheet():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "大表"
    for i in range(2100):
        ws.append([f"数据{i}", i])
    buf = io.BytesIO()
    wb.save(buf)
    r = parse_attachment("大表.xlsx", buf.getvalue())
    assert r.parse_status == "ok"
    assert sum(1 for ln in r.text.splitlines() if ln.startswith("数据")) == 2000
    assert any("截断" in n for n in r.notes)


def test_parse_xlsx_corrupt_returns_parse_failed():
    r = parse_attachment("坏.xlsx", b"PK\x03\x04 not-a-real-zip")
    assert r.fmt == "xlsx"
    assert r.parse_status == "parse_failed"


# ---------- 不支持 / 未知格式 ----------


@pytest.mark.parametrize(
    "filename,content,fmt",
    [
        ("表格.xls", OLE2_HEAD, "xls"),
        ("幻灯片.ppt", OLE2_HEAD, "ppt"),
        ("幻灯片.pptx", b"PK\x03\x04 junk", "pptx"),
        ("说明.txt", b"plain text", "unknown"),
    ],
)
def test_unsupported_formats_return_status_without_content(filename, content, fmt):
    r = parse_attachment(filename, content)
    assert r.fmt == fmt
    assert r.parse_status == "unsupported_format"
    assert r.text == ""
    assert r.char_count == 0
