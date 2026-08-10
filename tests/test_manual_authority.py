from __future__ import annotations

import json
from pathlib import Path
from urllib.parse import quote

import pytest
import yaml
from docx import Document

from sufe_qa.ingest.manual_authority import import_manual_authority_files
from sufe_qa.schema import (
    DocMeta,
    append_manifest,
    doc_id_from,
    load_manifest,
    load_relations,
    sha256_text,
)


def _write_docx(path: Path, title: str, body: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    document.save(path)


def _write_rules(path: Path, entries: list[dict[str, str]]) -> None:
    path.write_text(
        yaml.safe_dump(
            {"schema_version": "1", "namespace": "sufe-regulations", "entries": entries},
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )


def _entry(relative_path: str, **overrides: str) -> dict[str, str]:
    entry = {
        "path": relative_path,
        "category": "学工事务",
        "publisher": "上海财经大学教务处",
        "scope_unit": "上海财经大学",
        "source_section": "本科教学",
        "document_kind": "policy",
        "retention_status": "active",
    }
    entry.update(overrides)
    return entry


def test_dry_run_is_recursive_and_writes_nothing(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    accepted = source / "本科教学/20250801关于印发《课程考核管理办法》的通知.docx"
    draft = source / "本科教学/课程考核管理办法（征求意见稿）.docx"
    unlisted = source / "党建/内部管理办法.docx"
    broken = source / "本科教学/损坏文件.pdf"
    body = "第一条 为规范本科学生课程考核，明确申请条件、办理流程和审批部门。" * 4
    _write_docx(accepted, "课程考核管理办法", body)
    _write_docx(draft, "课程考核管理办法征求意见稿", body)
    _write_docx(unlisted, "内部管理办法", body)
    broken.parent.mkdir(parents=True, exist_ok=True)
    broken.write_bytes(b"not-a-pdf")
    rules = tmp_path / "rules.yaml"
    _write_rules(
        rules,
        [
            _entry(accepted.relative_to(source).as_posix()),
            _entry(draft.relative_to(source).as_posix()),
            _entry(broken.relative_to(source).as_posix()),
        ],
    )
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"

    report = import_manual_authority_files(
        source, rules, corpus, manifest, apply=False, evaluated_at="2026-08-10T00:00:00Z"
    )

    assert report.total_files == 4
    assert report.accepted == 1
    assert report.persisted == 0
    assert report.excluded == 2
    assert report.incomplete == 1
    assert report.quarantined == 0
    assert {decision.relative_path for decision in report.decisions} == {
        accepted.relative_to(source).as_posix(),
        draft.relative_to(source).as_posix(),
        unlisted.relative_to(source).as_posix(),
        broken.relative_to(source).as_posix(),
    }
    reasons = {decision.relative_path: decision.reason for decision in report.decisions}
    assert reasons[draft.relative_to(source).as_posix()] == "draft_filename"
    assert reasons[unlisted.relative_to(source).as_posix()] == "not_allowlisted"
    assert reasons[broken.relative_to(source).as_posix()].startswith("parse_")
    assert not manifest.exists()
    assert list(corpus.rglob("*.md")) == []


def test_apply_persists_stable_manual_metadata_and_deduplicates(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/20250801关于印发《课程考核管理办法》的通知.docx"
    document_path = source / relative
    _write_docx(
        document_path,
        "课程考核管理办法",
        "第一条 本办法适用于本科学生。第二条 学生应按规定参加课程考核。" * 5,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"

    first = import_manual_authority_files(
        source, rules, corpus, manifest, apply=True, evaluated_at="2026-08-10T00:00:00Z"
    )
    second = import_manual_authority_files(
        source, rules, corpus, manifest, apply=True, evaluated_at="2026-08-10T00:00:00Z"
    )

    assert first.accepted == first.persisted == 1
    assert second.accepted == 0
    assert second.duplicates == 1
    loaded = load_manifest(manifest)
    assert len(loaded) == 1
    meta = next(iter(loaded.values()))
    assert meta.source_url.startswith("manual://sufe-regulations/")
    assert "%E6%9C%AC%E7%A7%91%E6%95%99%E5%AD%A6" in meta.source_url
    assert meta.source_type == "manual_upload"
    assert meta.document_type == "attachment"
    assert meta.document_kind == "policy"
    assert meta.retention_status == "active"
    assert meta.index_collection == "main_qa"
    assert meta.validity_status == "unknown_validity"
    assert meta.publish_date == "2025-08-01"
    assert meta.publish_date_evidence == "文件名前缀：20250801"
    assert meta.binary_hash and meta.text_hash
    content = (corpus / meta.file_path).read_text(encoding="utf-8")
    assert meta.content_hash == sha256_text(content)


def test_rule_validation_rejects_duplicate_and_traversal_paths(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    source.mkdir()
    rules = tmp_path / "rules.yaml"
    duplicate = _entry("本科教学/办法.pdf")
    _write_rules(rules, [duplicate, duplicate])

    with pytest.raises(ValueError, match="重复规则路径"):
        import_manual_authority_files(
            source, rules, tmp_path / "corpus", tmp_path / "corpus/manifest.jsonl"
        )

    _write_rules(rules, [_entry("../办法.pdf")])
    with pytest.raises(ValueError, match="非法规则路径"):
        import_manual_authority_files(
            source, rules, tmp_path / "corpus", tmp_path / "corpus/manifest.jsonl"
        )


def test_apply_refreshes_explicit_metadata_for_same_source_and_content(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "研究生教学/旧版学位管理办法.docx"
    _write_docx(
        source / relative,
        "旧版学位管理办法",
        "第一条 本办法规定研究生申请学位的条件、材料和审核程序。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"
    _write_rules(rules, [_entry(relative, retention_status="active")])
    import_manual_authority_files(source, rules, corpus, manifest, apply=True)
    _write_rules(rules, [_entry(relative, retention_status="historical")])

    report = import_manual_authority_files(source, rules, corpus, manifest, apply=True)

    assert report.accepted == report.persisted == 1
    assert report.duplicates == 0
    meta = next(iter(load_manifest(manifest).values()))
    assert meta.retention_status == "historical"
    assert meta.index_collection == "historical"


def test_report_serializes_all_decisions(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    _write_docx(source / "未收录.docx", "未收录", "第一条 普通正文。" * 10)
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [])
    report_path = tmp_path / "report.json"

    report = import_manual_authority_files(
        source,
        rules,
        tmp_path / "corpus",
        tmp_path / "corpus/manifest.jsonl",
        report_path=report_path,
    )

    data = json.loads(report_path.read_text(encoding="utf-8"))
    assert data["total_files"] == 1
    assert data["decisions"][0]["status"] == "excluded"
    assert data["decisions"][0]["reason"] == "not_allowlisted"
    assert report.excluded == 1


def test_allowlisted_source_symlink_is_rejected_without_reading_target(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/课程考核管理办法.docx"
    outside = tmp_path / "outside.docx"
    _write_docx(outside, "外部文件", "第一条 外部文件不应被导入。" * 10)
    link = source / relative
    link.parent.mkdir(parents=True)
    link.symlink_to(outside)
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])

    report = import_manual_authority_files(
        source, rules, tmp_path / "corpus", tmp_path / "corpus/manifest.jsonl"
    )

    assert report.accepted == 0
    decision = next(item for item in report.decisions if item.relative_path == relative)
    assert decision.status == "excluded"
    assert decision.reason == "unsafe_source_symlink"


def test_unsafe_existing_manifest_path_cannot_escape_corpus(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/课程考核管理办法.docx"
    _write_docx(
        source / relative,
        "课程考核管理办法",
        "第一条 本办法适用于本科学生课程考核，明确办理流程和审批要求。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"
    source_url = f"manual://sufe-regulations/{quote(relative, safe='/')}"
    append_manifest(
        manifest,
        [
            DocMeta(
                doc_id=doc_id_from(source_url),
                title="课程考核管理办法",
                source_url=source_url,
                publisher="上海财经大学教务处",
                publish_date="unknown",
                category="学工事务",
                fetched_at="2026-08-10T00:00:00Z",
                content_hash="sha256:unsafe",
                file_path="../escaped.md",
                quality_status="accepted",
                parse_status="ok",
                document_kind="policy",
                source_type="manual_upload",
                source_section="本科教学",
                scope_unit="上海财经大学",
                retention_status="active",
                index_collection="main_qa",
            )
        ],
    )
    escaped = tmp_path / "escaped.md"
    escaped.write_text("guard", encoding="utf-8")

    report = import_manual_authority_files(source, rules, corpus, manifest, apply=True)

    assert escaped.read_text(encoding="utf-8") == "guard"
    assert report.persisted == 0
    decision = next(item for item in report.decisions if item.relative_path == relative)
    assert decision.status == "quarantined"
    assert decision.reason == "unsafe_existing_file_path"


def test_hard_subject_and_declared_kind_conflicts_are_excluded(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    internal = "内部/党委经费报销管理办法.docx"
    news = "本科教学/学院新闻活动回顾.docx"
    _write_docx(
        source / internal,
        "党委经费报销管理办法",
        "第一条 本办法规定党务经费和教职工差旅报销审批流程。" * 6,
    )
    _write_docx(
        source / news,
        "学院新闻活动回顾",
        "学院组织师生参加宣传活动，召开交流会议并回顾活动成果。" * 8,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(internal), _entry(news)])

    report = import_manual_authority_files(
        source, rules, tmp_path / "corpus", tmp_path / "corpus/manifest.jsonl"
    )

    reasons = {item.relative_path: item.reason for item in report.decisions}
    assert reasons[internal] == "hard_excluded_subject"
    assert reasons[news] == "declared_kind_conflict:news"
    assert report.accepted == 0


def test_hard_subject_gate_does_not_reject_incidental_body_references(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/关于印发《本科学生学籍管理实施细则》的通知.docx"
    _write_docx(
        source / relative,
        "关于印发《本科学生学籍管理实施细则》的通知",
        "本细则已经学校党委会议审议通过。第一条 本细则明确学生学籍办理流程。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])

    report = import_manual_authority_files(
        source, rules, tmp_path / "corpus", tmp_path / "corpus/manifest.jsonl"
    )

    assert report.accepted == 1
    assert report.excluded == 0


def test_removed_allowlist_entry_revokes_materialized_document(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/课程考核管理办法.docx"
    _write_docx(
        source / relative,
        "课程考核管理办法",
        "第一条 本办法适用于本科学生课程考核，明确办理流程和审批要求。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"
    _write_rules(rules, [_entry(relative)])
    import_manual_authority_files(source, rules, corpus, manifest, apply=True)
    before = next(iter(load_manifest(manifest).values()))
    old_file = corpus / before.file_path
    assert old_file.is_file()
    _write_rules(rules, [])

    report = import_manual_authority_files(source, rules, corpus, manifest, apply=True)

    assert report.revoked == 1
    assert not old_file.exists()
    after = load_manifest(manifest)[before.doc_id]
    assert after.retention_status == "archived"
    assert after.retention_reason == "manual_allowlist_revoked"
    assert after.index_collection == "none"
    assert after.file_path == ""
    assert after.content_hash == ""


def test_rejected_hash_record_does_not_block_authority_import(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/课程考核管理办法.docx"
    _write_docx(
        source / relative,
        "课程考核管理办法",
        "第一条 本办法适用于本科学生课程考核，明确办理流程和审批要求。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"
    probe = import_manual_authority_files(source, rules, corpus, manifest)
    text_hash = next(item.text_hash for item in probe.decisions if item.status == "accepted")
    append_manifest(
        manifest,
        [
            DocMeta(
                doc_id=doc_id_from("https://example.sufe.edu.cn/rejected"),
                title="旧空壳",
                source_url="https://example.sufe.edu.cn/rejected",
                publisher="上海财经大学",
                publish_date="unknown",
                category="学工事务",
                fetched_at="2026-08-10T00:00:00Z",
                content_hash="",
                text_hash=text_hash,
                file_path="",
                parse_status="ok",
                quality_status="low_quality",
            )
        ],
    )

    report = import_manual_authority_files(source, rules, corpus, manifest, apply=True)

    assert report.accepted == report.persisted == 1
    assert report.duplicates == 0


def test_cross_source_duplicate_keeps_manual_alias_and_relation(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/课程考核管理办法.docx"
    _write_docx(
        source / relative,
        "课程考核管理办法",
        "第一条 本办法适用于本科学生课程考核，明确办理流程和审批要求。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"
    probe = import_manual_authority_files(source, rules, corpus, manifest)
    manual_decision = next(item for item in probe.decisions if item.status == "accepted")
    existing_url = "https://jwc.sufe.edu.cn/policy/course-assessment"
    existing_id = doc_id_from(existing_url)
    existing_rel = "学工事务/existing.md"
    existing_text = "# 已抓取课程考核管理办法\n\n已有权威正文。\n"
    (corpus / existing_rel).parent.mkdir(parents=True, exist_ok=True)
    (corpus / existing_rel).write_text(existing_text, encoding="utf-8")
    append_manifest(
        manifest,
        [
            DocMeta(
                doc_id=existing_id,
                title="已抓取课程考核管理办法",
                source_url=existing_url,
                publisher="上海财经大学教务处",
                publish_date="2025-08-01",
                category="学工事务",
                fetched_at="2026-08-10T00:00:00Z",
                content_hash=sha256_text(existing_text),
                text_hash=manual_decision.text_hash,
                file_path=existing_rel,
                parse_status="ok",
                quality_status="accepted",
                document_kind="policy",
                source_type="official_department",
                source_section="规章制度",
                scope_unit="本科生",
                retention_status="active",
                index_collection="main_qa",
            )
        ],
    )

    report = import_manual_authority_files(source, rules, corpus, manifest, apply=True)

    assert report.duplicates == 1
    manual_meta = load_manifest(manifest)[manual_decision.doc_id]
    assert manual_meta.quality_status == "duplicate"
    assert manual_meta.canonical_doc_id == existing_id
    relation = next(
        item
        for item in load_relations(corpus / "relations.jsonl")
        if item.relation == "same_content_as"
    )
    assert relation.parent_doc_id == manual_decision.doc_id
    assert relation.child_doc_id == existing_id


def test_same_text_changed_binary_refreshes_binary_hash(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    relative = "本科教学/课程考核管理办法.docx"
    path = source / relative
    _write_docx(
        path,
        "课程考核管理办法",
        "第一条 本办法适用于本科学生课程考核，明确办理流程和审批要求。" * 6,
    )
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])
    corpus = tmp_path / "corpus"
    manifest = corpus / "manifest.jsonl"
    import_manual_authority_files(source, rules, corpus, manifest, apply=True)
    before = next(iter(load_manifest(manifest).values()))
    document = Document(path)
    document.core_properties.comments = "重新签章导出"
    document.save(path)

    report = import_manual_authority_files(source, rules, corpus, manifest, apply=True)

    after = load_manifest(manifest)[before.doc_id]
    assert report.accepted == report.persisted == 1
    assert after.text_hash == before.text_hash
    assert after.binary_hash != before.binary_hash


def test_missing_allowlisted_file_is_reported(tmp_path: Path) -> None:
    source = tmp_path / "规章制度"
    source.mkdir()
    relative = "本科教学/缺失的课程管理办法.pdf"
    rules = tmp_path / "rules.yaml"
    _write_rules(rules, [_entry(relative)])

    report = import_manual_authority_files(
        source, rules, tmp_path / "corpus", tmp_path / "corpus/manifest.jsonl"
    )

    assert report.missing == 1
    decision = next(item for item in report.decisions if item.relative_path == relative)
    assert decision.status == "missing"
    assert decision.reason == "allowlisted_file_missing"
