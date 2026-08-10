"""CLI 离线端到端：ingest → index → ask → eval → crawl 胶水，全程 FakeEmbedder/FakeLLM。"""

from __future__ import annotations

import json

import pytest
import yaml
from docx import Document

from sufe_qa import cli
from sufe_qa.config import load_settings
from sufe_qa.coverage.question_bank import QuestionBank, QuestionProbe
from sufe_qa.generate.client import FakeLLM
from sufe_qa.retrieve.retriever import Hit
from sufe_qa.schema import doc_id_from, load_manifest

DOC_TEXT = (
    "推免工作实施办法 第一条 申请推免的学生应为纳入国家普通本科招生计划录取的应届毕业生，"
    "拥护中国共产党的领导，品德良好，遵纪守法，身心健康，诚实守信，学风端正。"
    "第二条 申请学生应勤奋学习，刻苦钻研，成绩优秀，学术研究兴趣浓厚。"
) * 2
QUESTION = DOC_TEXT[10:110]  # 用正文子串提问：FakeEmbedder 下保证高相似度过门控


@pytest.fixture
def settings(tmp_path, monkeypatch):
    monkeypatch.setenv("SUFE_QA_DATA_DIR", str(tmp_path))
    monkeypatch.setattr(cli, "_make_llm", lambda s: FakeLLM(1))
    s = load_settings()
    (s.inbox_dir / "tuimian.md").write_text("# 推免工作实施办法\n\n" + DOC_TEXT, encoding="utf-8")
    return s


def test_help_exits_zero(capsys):
    with pytest.raises(SystemExit) as e:
        cli.main(["--help"])
    assert e.value.code == 0


def test_ingest_authority_files_cli_defaults_to_dry_run(settings, tmp_path, capsys):
    source = tmp_path / "规章制度"
    relative = "本科教学/20250801课程考核管理办法.docx"
    path = source / relative
    path.parent.mkdir(parents=True)
    document = Document()
    document.add_heading("课程考核管理办法", level=1)
    document.add_paragraph("第一条 本办法适用于本科学生课程考核和成绩管理。" * 8)
    document.save(path)
    rules = tmp_path / "rules.yaml"
    rules.write_text(
        yaml.safe_dump(
            {
                "schema_version": "1",
                "namespace": "sufe-regulations",
                "entries": [
                    {
                        "path": relative,
                        "category": "学工事务",
                        "publisher": "上海财经大学教务处",
                        "scope_unit": "上海财经大学",
                        "source_section": "本科教学",
                        "document_kind": "policy",
                        "retention_status": "active",
                    }
                ],
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    report = tmp_path / "report.json"

    assert (
        cli.main(
            [
                "ingest-authority-files",
                "--source",
                str(source),
                "--rules",
                str(rules),
                "--report",
                str(report),
            ]
        )
        == 0
    )
    assert load_manifest(settings.manifest_path) == {}
    assert "可导入 1，实际写入 0" in capsys.readouterr().out

    assert (
        cli.main(
            [
                "ingest-authority-files",
                "--source",
                str(source),
                "--rules",
                str(rules),
                "--report",
                str(report),
                "--apply",
            ]
        )
        == 0
    )
    assert len(load_manifest(settings.manifest_path)) == 1


def test_coverage_audit_writes_json_and_markdown(settings, tmp_path):
    json_path = tmp_path / "coverage.json"
    markdown_path = tmp_path / "coverage.md"
    assert (
        cli.main(
            [
                "coverage-audit",
                "--question-bank",
                "data/eval/sufe_question_bank.jsonl",
                "--manifest",
                str(settings.manifest_path),
                "--corpus",
                str(settings.corpus_dir),
                "--output-json",
                str(json_path),
                "--output-md",
                str(markdown_path),
            ]
        )
        == 0
    )
    assert json_path.is_file()
    assert markdown_path.is_file()
    assert "本科教务" in markdown_path.read_text(encoding="utf-8")


def test_answer_benchmark_cli_writes_real_answer_and_resumes(
    settings, tmp_path, monkeypatch, capsys
):
    probe = QuestionProbe(
        id="jwc-leave-001",
        question="本科生如何申请缓考？",
        scene="本科教务",
        required_source_type="official_procedure",
        expected_domains=("jwc.sufe.edu.cn",),
        expected_doc_ids=(),
        required_answer_points=("申请条件",),
        needs_current_version=True,
    )
    bank = QuestionBank((probe,), content_hash="sha256:bank")
    hit = Hit(
        chunk_id="doc-1::0000",
        doc_id="doc-1",
        title="缓考办理办法",
        category="本科教务",
        source_url="https://jwc.sufe.edu.cn/page.htm",
        publisher="上海财经大学教务处",
        heading_path="缓考",
        text="因病不能考试的学生应提交申请。",
        rrf_score=0.03,
        vector_similarity=0.9,
    )

    class StubRetriever:
        calls = 0

        def __init__(self, *args, **kwargs):
            pass

        def search_routed(self, question):
            type(self).calls += 1
            return [hit]

    monkeypatch.setattr(cli, "load_question_bank", lambda path: bank)
    monkeypatch.setattr(
        cli,
        "load_index_metadata",
        lambda settings: {
            "index_fingerprint": "sha256:index",
            "embedding_model": "fake-for-test",
            "embedding_backend": "fake",
            "test_only": True,
        },
    )
    monkeypatch.setattr(cli, "HybridRetriever", StubRetriever)
    output = tmp_path / "real_answers.json"

    argv = [
        "answer-benchmark",
        "--bank",
        "ignored.jsonl",
        "--output-json",
        str(output),
        "--workers",
        "1",
        "--fake-embed",
    ]
    assert cli.main(argv) == 0
    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["results"][0]["status"] == "answered"
    assert "离线演示回答" in payload["results"][0]["answer_text"]
    assert "[1/1] answered" in capsys.readouterr().out

    StubRetriever.calls = 0
    assert cli.main([*argv, "--resume"]) == 0
    assert StubRetriever.calls == 0


def test_quality_audit_and_clean_rebuild_preview_do_not_mutate_corpus(settings, tmp_path):
    assert cli.main(["ingest", "--category", "学工事务", "--publisher", "研究生院"]) == 0
    json_path = tmp_path / "quality.json"
    markdown_path = tmp_path / "quality.md"
    assert (
        cli.main(
            [
                "quality-audit",
                "--manifest",
                str(settings.manifest_path),
                "--corpus",
                str(settings.corpus_dir),
                "--raw",
                str(settings.data_dir / "raw"),
                "--output-json",
                str(json_path),
                "--output-md",
                str(markdown_path),
            ]
        )
        == 0
    )
    before = settings.manifest_path.read_bytes()
    assert cli.main(["rebuild-clean-corpus", "--audit", str(json_path)]) == 0
    assert settings.manifest_path.read_bytes() == before
    assert json_path.is_file() and markdown_path.is_file()


def test_offline_end_to_end(settings, capsys):
    assert cli.main(["ingest", "--category", "学工事务", "--publisher", "研究生院"]) == 0
    assert cli.main(["index", "--fake-embed"]) == 0

    assert cli.main(["ask", QUESTION, "--fake-embed"]) == 0
    out = capsys.readouterr().out
    assert "来源：" in out
    assert "推免工作实施办法" in out

    # 无稽问题：相似度过不了门控 → 拒答模板，无来源卡片
    assert cli.main(["ask", "qwerty asdfg zxcvb", "--fake-embed"]) == 0
    out = capsys.readouterr().out
    assert "未在已收录的学校官方资料中找到可靠依据" in out
    assert "来源：" not in out


def test_eval_command_gate(settings, tmp_path, capsys):
    cli.main(["ingest", "--category", "学工事务"])
    cli.main(["index", "--fake-embed"])
    evalset = tmp_path / "evalset.jsonl"
    evalset.write_text(
        '{"id": "q1", "question": "%s", "expected_doc_ids": ["%s"]}\n'
        '{"id": "q2", "question": "qwerty asdfg", "should_refuse": true}\n'
        % (QUESTION.replace('"', "'"), doc_id_from("inbox/tuimian.md")),
        encoding="utf-8",
    )
    assert cli.main(["eval", "--evalset", str(evalset), "--fake-embed"]) == 0
    out = capsys.readouterr().out
    assert "检索命中率: 100.0%" in out
    assert "拒答正确率: 100.0%" in out
    # 达标线拉满到不可能的水平 → 门禁失败，退出码 1
    capsys.readouterr()
    assert (
        cli.main(["eval", "--evalset", str(evalset), "--fake-embed", "--min-refusal", "1.1"]) == 1
    )


def test_crawl_glues_pages_into_corpus(settings, tmp_path, monkeypatch):
    """crawl 走新引擎：stub SafeFetcher 返回列表页+文章页，验证入库与 doc_id 锚定。"""
    from sufe_qa.crawler.fetcher import FetchResult

    article_url = "https://gs.sufe.edu.cn/Home/Detail/8001"
    list_url = "https://gs.sufe.edu.cn/Home/List/31"
    routes = {
        list_url: FetchResult(
            requested_url=list_url,
            final_url=list_url,
            status_code=200,
            content=f'<html><body><div class="blog-content"><a href="{article_url}">推免通知</a></div></body></html>'.encode(),
        ),
        article_url: FetchResult(
            requested_url=article_url,
            final_url=article_url,
            status_code=200,
            content=f"<html><body><article><h1>推免通知</h1><p>{DOC_TEXT}</p></article></body></html>".encode(),
        ),
    }

    class StubFetcher:
        def __init__(self, *a, **kw):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *exc):
            pass

        def fetch(self, url, kind="html", headers=None):
            return routes.get(url) or FetchResult(
                requested_url=url, final_url=url, status="http_error", status_code=404, error="404"
            )

    monkeypatch.setattr(cli, "SafeFetcher", StubFetcher)
    seeds = tmp_path / "seeds.yaml"
    seeds.write_text(
        "seeds:\n"
        "  - name: 测试种子\n"
        f"    list_url: {list_url}\n"
        '    link_selector: ".blog-content a"\n'
        "    url_prefix: https://gs.sufe.edu.cn/Home/Detail/\n"
        "    category: 推免升学\n"
        "    publisher: 研究生院\n",
        encoding="utf-8",
    )
    assert cli.main(["crawl", "--seeds", str(seeds)]) == 0
    manifest = load_manifest(settings.manifest_path)
    assert len(manifest) == 1
    meta = next(iter(manifest.values()))
    assert meta.source_url == article_url
    assert meta.doc_id == doc_id_from(article_url)
    assert meta.category == "推免升学"
    assert meta.quality_status == "accepted"
    assert meta.document_type == "article"
    # 抓取状态已持久化，raw 缓存已落地
    assert (settings.data_dir / "crawl_state" / "gs.sufe.edu.cn.json").is_file()
    assert (
        settings.data_dir / "raw" / "gs.sufe.edu.cn" / "articles" / f"{meta.doc_id}.html"
    ).is_file()
