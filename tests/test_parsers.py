import pytest

from sufe_qa.ingest.parsers import _strip_nav_soup, parse_file, parse_html


def test_strip_nav_soup_removes_giant_nav_line():
    giant = (
        "上财首页 English 首页 教职员工 师资队伍 专职辅导员 行政人员 特聘教授 "
        "师资风采 科学研究 科研成果 科研项目 科研机构 讲座会议预告 期刊目录 本科生培养"
    )
    body = "经济学院2025年博士研究生综合考核方案\n第一条 坚持公平公正原则。"
    assert _strip_nav_soup(f"{giant}\n{body}") == body
    # 无站点标志词的短词大行不剥（可能是正文列举）
    not_nav = "苹果 香蕉 橙子 西瓜 葡萄 桃子 李子 杏子 樱桃 草莓"
    assert _strip_nav_soup(f"{not_nav}\n{body}").startswith(not_nav)


def test_strip_nav_soup_removes_boilerplate():
    nav_head = "\n".join(["邮箱", "信息门户", "上财首页", "English", "师资队伍", "科学研究"])
    nav_tail = "\n".join(["上一篇", "扫码关注我们", "通知公告"])
    body = "第一条 为规范研究生招生工作，制定本规定。\n第二条 招生工作应遵循公平原则。"
    assert _strip_nav_soup(f"{nav_head}\n{body}\n{nav_tail}") == body


def test_strip_nav_soup_keeps_short_content_lines():
    # 正文开头的偶发短行（不足连续 5 行）不剥
    text = "总则\n第一条 为规范招生工作，制定本规定。\n第二条 内容如下。"
    assert _strip_nav_soup(text) == text


def test_strip_nav_soup_removes_mid_body_nav_block():
    # gs 站面包屑"首页 / 管理规定"打断头尾连续块后，中段残留的子菜单导航块
    mid_nav = "\n".join(
        [
            "招生信息",
            "通知公告",
            "硕士生招生",
            "博士生招生",
            "港澳台招生",
            "招生简章",
            "管理规定",
            "公示专栏",
            "历史数据",
            "硕士生招生",
            "博士生招生",
            "2026硕士招生专题",
        ]
    )
    text = (
        f"上海财经大学研究生院\n首页\n/ 管理规定\n{mid_nav}\n"
        "关于印发《XX办法》的通知\n发布时间：2025-07-28"
    )
    out = _strip_nav_soup(text)
    assert "通知公告" not in out
    assert "硕士生招生" not in out
    assert "关于印发《XX办法》的通知" in out
    assert "发布时间：2025-07-28" in out
    # 中段无机构菜单词汇的连续短行块（如录取名单）不剥
    names = "\n".join(["张三", "李四", "王五", "赵六", "陈七", "林八", "黄九", "周十"])
    name_text = f"录取名单如下：\n{names}\n以上名单公示五天。"
    assert _strip_nav_soup(name_text) == name_text


def test_parse_html_strips_nav_but_keeps_body():
    nav = "".join(f"<li>{w}</li>" for w in ["邮箱", "信息门户", "上财首页", "English", "师资队伍"])
    html = (
        f"<html><head><title>考核方案</title></head><body>"
        f"<ul>{nav}</ul><article><p>经济学院2025年博士研究生综合考核方案如下。</p>"
        f"<p>第一条 坚持公平公正。</p></article></body></html>"
    )
    doc = parse_html(html, "fallback")
    assert "邮箱" not in doc.text
    assert "考核方案" in doc.text


def test_parse_html_breadcrumb_title_first_segment():
    # 高校 CMS 面包屑标题"文章名|栏目|… - 站名"，取首段为文章名
    html = (
        "<html><head><title>关于印发《研究生招生工作管理规定》的通知"
        "|管理规定|招生信息 - 上海财经大学研究生院</title></head>"
        "<body><article><p>第一条 为规范研究生招生工作，制定本规定。</p></article></body></html>"
    )
    doc = parse_html(html, "fallback")
    assert doc.title == "关于印发《研究生招生工作管理规定》的通知"


def test_parse_html_extracts_text_and_title():
    html = """<html><head><title>关于开展2025年奖学金评审的通知</title></head>
    <body><article><p>各学院：现将评审安排通知如下。</p><p>申请条件如下。</p></article></body></html>"""
    doc = parse_html(html, "fallback")
    assert "申请条件" in doc.text
    assert "奖学金" in doc.title


def test_parse_pdf(tmp_path):
    import fitz

    p = tmp_path / "rule.pdf"
    d = fitz.open()
    page = d.new_page()
    # 默认 helv 字体不含 CJK 字形，需指定内置中文字体，否则写入的是乱码
    page.insert_text((72, 72), "第一条 本办法适用于全日制在校生。", fontname="china-s")
    d.save(str(p))
    d.close()
    doc = parse_file(p)
    assert "第一条" in doc.text
    assert doc.title == "rule"


def test_parse_docx(tmp_path):
    from docx import Document

    p = tmp_path / "notice.docx"
    d = Document()
    d.add_paragraph("推免工作将于九月启动。")
    d.save(str(p))
    doc = parse_file(p)
    assert "推免" in doc.text


def test_parse_md_passthrough(tmp_path):
    p = tmp_path / "note.md"
    p.write_text("# 学工指南\n内容正文。", encoding="utf-8")
    doc = parse_file(p)
    assert "内容正文" in doc.text  # H1 已提为 title，正文不再含标题行
    assert doc.title == "学工指南"  # md 以首个一级标题为文档标题


def test_unsupported_suffix(tmp_path):
    p = tmp_path / "a.exe"
    p.write_bytes(b"x")
    with pytest.raises(ValueError, match="不支持"):
        parse_file(p)


def test_parse_html_gbk_fallback(tmp_path):
    html = "<html><head><title>奖学金通知</title></head><body><p>申请条件如下</p></body></html>"
    p = tmp_path / "gbk.html"
    p.write_bytes(html.encode("gb18030"))
    doc = parse_file(p)
    assert "申请条件" in doc.text


def test_parse_docx_includes_tables(tmp_path):
    from docx import Document

    p = tmp_path / "table.docx"
    d = Document()
    d.add_paragraph("标准如下：")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "国家奖学金"
    t.rows[0].cells[1].text = "8000元"
    d.save(str(p))
    doc = parse_file(p)
    assert "8000元" in doc.text


def test_corrupt_pdf_raises_parse_error(tmp_path):
    from sufe_qa.ingest.parsers import ParseError

    p = tmp_path / "bad.pdf"
    p.write_bytes(b"not a pdf at all")
    with pytest.raises(ParseError):
        parse_file(p)
