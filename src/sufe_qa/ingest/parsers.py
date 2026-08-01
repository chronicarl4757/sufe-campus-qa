from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class ParsedDoc:
    title: str
    text: str
    publish_date: str = "unknown"
    publisher: str = ""


class ParseError(Exception):
    """解析失败的统一异常，携带文件路径与原始原因。"""


def _read_html_text(path: Path) -> str:
    """读取 html 文件，兼容高校老网站的 GBK/GB2312 编码（gb18030 是其超集）。"""
    raw = path.read_bytes()
    for enc in ("utf-8", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


_TITLE_RE = re.compile(r"<title[^>]*>(.*?)</title>", re.S | re.I)

# 导航样板行：短且无句读/括号（菜单项 "校友风采/暑期师资班/研究生暑期学校" 长短不一）
_NAV_LINE = re.compile(r"^[^\s。，；：、！？.,;:!?()（）《》【】\[\]{}]{1,15}$")
# 机构菜单词汇：导航块判定的必要证据；正文/名单不含此类词，防止误剥
_NAV_MARKER = re.compile(
    r"首页|English|信息门户|邮箱|师资队伍|导航|关注我们|扫码|校友|学院简介"
    r"|科学研究|本科生培养|研究生培养|组织机构|学院领导|党建工作|新闻中心|通知公告"
)


def _is_giant_nav_line(line: str) -> bool:
    """整行由空格分隔的短链接词组成且含站点标志词（_wp3 导航菜单的常见抽取形态）。"""
    toks = line.split()
    if len(toks) < 8:
        return False
    short = sum(1 for t in toks if len(t) <= 6)
    return short / len(toks) >= 0.7 and bool(_NAV_MARKER.search(line))


def _nav_run(lines: list[str], min_len: int) -> int:
    """从头量出导航样板块长度：短无标点行连续 >=min_len 且块内含机构菜单词汇。"""
    run = 0
    for line in lines:
        if _NAV_LINE.match(line.strip()):
            run += 1
        else:
            break
    if run >= min_len and any(_NAV_MARKER.search(line) for line in lines[:run]):
        return run
    return 0


def _strip_mid_runs(lines: list[str], min_len: int = 8) -> list[str]:
    """剥掉正文中段残留的导航行块：连续 >=min_len 短行且含机构菜单词汇。

    中段阈值高于头尾（面包屑行会打断头尾连续块，如 gs 站"首页 / 管理规定"），
    短名单/条款等正文达不到 8 行连续短行且含菜单词的双重条件。
    """
    out, i, n = [], 0, len(lines)
    while i < n:
        if not _NAV_LINE.match(lines[i].strip()):
            out.append(lines[i])
            i += 1
            continue
        j = i
        while j < n and _NAV_LINE.match(lines[j].strip()):
            j += 1
        if not (j - i >= min_len and any(_NAV_MARKER.search(ln) for ln in lines[i:j])):
            out.extend(lines[i:j])
        i = j
    return out


def _strip_nav_soup(text: str, head_min: int = 5, tail_min: int = 3) -> str:
    """剥掉 _wp3 站正文的导航样板。

    三种形态：①单行巨型导航行（头 4 行/尾 2 行内判定即剥）；
    ②连续多行短导航行（头 >=head_min 行、尾 >=tail_min 行且含机构菜单词汇才剥）；
    ③中段残留导航块（>=8 连续短行且含菜单词，如 gs 站面包屑后的子菜单块）。
    阈值与词汇证据双重条件，保证不误伤正文短行、条款列表与名单类正文。
    """
    lines = text.split("\n")
    kill = set()
    for idx in range(min(4, len(lines))):
        if _is_giant_nav_line(lines[idx].strip()):
            kill.add(idx)
    for idx in range(max(0, len(lines) - 2), len(lines)):
        if _is_giant_nav_line(lines[idx].strip()):
            kill.add(idx)
    lines = [ln for i, ln in enumerate(lines) if i not in kill]

    head = _nav_run(lines, head_min)
    if head:
        del lines[:head]
    tail = _nav_run(lines[::-1], tail_min)
    if tail:
        del lines[len(lines) - tail :]
    return "\n".join(_strip_mid_runs(lines))


def _breadcrumb_title(raw: str) -> str | None:
    """高校 CMS 常用面包屑标题"文章名|栏目|… - 站名"，首段才是文章名；
    只在 <title> 含分隔符时出手，其余情况交给 trafilatura。"""
    m = _TITLE_RE.search(raw)
    if not m:
        return None
    t = re.sub(r"\s+", " ", m.group(1)).strip()
    if "|" not in t and "｜" not in t:
        return None
    return re.split(r"[|｜]", t)[0].strip() or None


def parse_html(raw: str, fallback_title: str) -> ParsedDoc:
    import trafilatura

    text = trafilatura.extract(raw, include_comments=False, include_tables=True) or ""
    meta = trafilatura.extract_metadata(raw)
    title = fallback_title
    if meta and meta.title:
        # 空白标题不覆盖 fallback
        title = meta.title.strip() or fallback_title
    title = _breadcrumb_title(raw) or title
    # trafilatura 不同版本的 metadata 可能没有 date 属性或为 None，统一兜底
    date = getattr(meta, "date", None) or "unknown"
    return ParsedDoc(title=title, text=_strip_nav_soup(text.strip()), publish_date=date)


def parse_pdf(path: Path) -> ParsedDoc:
    """解析 PDF。契约：扫描件/加密 PDF 返回空 text，由下游按空文档处理，不视为解析错误。"""
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    try:
        text = "\n".join(page.get_text().strip() for page in doc)
    finally:
        doc.close()
    return ParsedDoc(title=path.stem, text=text.strip())


def parse_docx(path: Path) -> ParsedDoc:
    from docx import Document

    d = Document(str(path))
    text_parts = [p.text for p in d.paragraphs if p.text.strip()]
    # 政策文件常用表格承载关键信息（金额、条件），需一并抽取
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))
    return ParsedDoc(title=path.stem, text="\n".join(text_parts).strip())


def parse_file(path: Path) -> ParsedDoc:
    """按后缀分发解析。

    不支持的后缀抛 ValueError；文件损坏等脏数据统一包装为 ParseError。
    下游 inbox 对两者都按错误计数跳过。
    """
    suffix = path.suffix.lower()
    if suffix not in (".html", ".htm", ".pdf", ".docx", ".md"):
        raise ValueError(f"不支持的文件类型: {path.name}")
    try:
        if suffix in (".html", ".htm"):
            return parse_html(_read_html_text(path), path.stem)
        if suffix == ".pdf":
            return parse_pdf(path)
        if suffix == ".docx":
            return parse_docx(path)
        text = path.read_text(encoding="utf-8", errors="replace").strip()
        # md 以首个一级标题为文档标题，并从正文剔除该行（避免落盘后双 H1）；
        # 无 H1 时回退文件名、正文原样
        title = path.stem
        m = re.search(r"(?m)^#\s+(.+)$", text)
        if m:
            title = m.group(1).strip()
            text = re.sub(r"(?m)^#\s+.+\n?", "", text, count=1).strip()
        return ParsedDoc(title=title, text=text)
    except Exception as e:
        raise ParseError(f"解析失败: {path.name}: {e}") from e
