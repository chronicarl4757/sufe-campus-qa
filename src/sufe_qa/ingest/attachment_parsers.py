"""网页附件解析：从字节流提取 PDF/DOCX/XLSX 文本，统一收敛为 AttachmentParseResult。

与 parsers.py（本地文件路径解析）不同，本模块面向爬虫下载的内存字节流：
格式按扩展名 + 魔数双重识别（冲突时以魔数为准），任何解析异常都不抛出，
收敛为 parse_status="parse_failed" 并在 notes 记录原因，由下游按状态分流。
"""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path

# 魔数
_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"  # docx/xlsx/pptx 均为 zip 容器
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0"  # doc/xls/ppt 均为 OLE2 复合文档

# 扩展名 -> 格式（无可靠魔数时的兜底识别）
_KNOWN_EXTS = {"pdf", "docx", "doc", "xlsx", "xls", "ppt", "pptx"}

# PDF 扫描件判定：平均每页文本低于该字符数视为扫描件（不强制 OCR，仅标记）
_SCANNED_AVG_CHARS = 50
# XLSX 单表最大输出行数，超出截断并在 notes 说明
_MAX_SHEET_ROWS = 2000
# 超过该字节数的 XLSX 改用 read_only 模式兜底（放弃合并单元格填充）
_READ_ONLY_BYTES = 5 * 1024 * 1024
# LibreOffice 转换超时（秒）
_SOFFICE_TIMEOUT = 60


@dataclass
class AttachmentParseResult:
    filename: str
    fmt: str  # pdf | docx | doc | xlsx | xls | ppt | pptx | unknown
    text: str  # 提取文本（markdown 化）；未解析为空串
    char_count: int
    page_count: int | None  # PDF 页数
    sheet_count: int | None  # XLSX 工作表数
    parse_status: str  # ok | scanned_pdf | legacy_doc_unparsed | unsupported_format | parse_failed
    notes: list[str] = field(default_factory=list)


def _ooxml_kind(content: bytes) -> str | None:
    """窥探 zip 包内结构区分 docx/xlsx/pptx（三者魔数相同，扩展名不可信时用）。"""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            names = set(zf.namelist())
    except zipfile.BadZipFile:
        return None
    if "word/document.xml" in names:
        return "docx"
    if "xl/workbook.xml" in names:
        return "xlsx"
    if "ppt/presentation.xml" in names:
        return "pptx"
    return None


def _detect_format(filename: str, content: bytes) -> str:
    """扩展名 + 魔数双重识别；冲突时以魔数为准。"""
    ext = Path(filename).suffix.lower().lstrip(".")
    if content.startswith(_PDF_MAGIC):
        return "pdf"
    if content.startswith(_ZIP_MAGIC):
        # OOXML 三家魔数相同：优先按扩展名细分，扩展名不在其中则窥探包内结构
        if ext in ("docx", "xlsx", "pptx"):
            return ext
        return _ooxml_kind(content) or "unknown"
    if content.startswith(_OLE2_MAGIC):
        # OLE2 同理按扩展名细分；扩展名不在 legacy 三件套内则无法可靠区分
        return ext if ext in ("doc", "xls", "ppt") else "unknown"
    return ext if ext in _KNOWN_EXTS else "unknown"


def _extract_pdf(result: AttachmentParseResult, content: bytes) -> None:
    """逐页提取 PDF 文本；平均每页文本过少判定为扫描件（仅标记，不 OCR）。"""
    import fitz  # pymupdf

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        pages = [page.get_text().strip() for page in doc]
        result.page_count = doc.page_count
    finally:
        doc.close()
    result.text = "\n".join(p for p in pages if p).strip()
    if result.page_count:
        avg = sum(len(p) for p in pages) / result.page_count
        if avg < _SCANNED_AVG_CHARS:
            result.parse_status = "scanned_pdf"
            # 定位首个稀疏页，提示从哪页起疑似扫描（整本扫描则为第 1 页）
            first_sparse = next(
                (i + 1 for i, p in enumerate(pages) if len(p) < _SCANNED_AVG_CHARS), 1
            )
            result.notes.append(f"平均每页文本 {avg:.0f} 字符，第{first_sparse}页起疑似扫描件")


def _extract_docx_text(content: bytes) -> str:
    """提取 DOCX 段落与表格：标题行加 '## ' 前缀，表格行沿用 ' | ' 拼接风格。"""
    from docx import Document

    d = Document(io.BytesIO(content))
    parts = []
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # 标题段落（Heading N / 标题 N）markdown 化为二级标题
        if p.style is not None and re.match(r"^(Heading|标题)\s*\d+$", p.style.name or ""):
            parts.append(f"## {text}")
        else:
            parts.append(text)
    # 政策文件常用表格承载关键信息（金额、条件），需一并抽取
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_doc(result: AttachmentParseResult, content: bytes) -> None:
    """legacy .doc 适配器：有 LibreOffice 则转换为 docx 再解析，否则放弃不报错。"""
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        result.parse_status = "legacy_doc_unparsed"
        result.notes.append("未检测到 LibreOffice/soffice，跳过 .doc 内容解析")
        return
    with tempfile.TemporaryDirectory(prefix="sufe_doc_") as tmp:
        src = Path(tmp) / "input.doc"
        src.write_bytes(content)
        try:
            proc = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(src)],
                timeout=_SOFFICE_TIMEOUT,
                capture_output=True,
            )
        except (subprocess.TimeoutExpired, OSError) as e:
            result.parse_status = "legacy_doc_unparsed"
            result.notes.append(f"LibreOffice 转换失败: {e}")
            return
        converted = src.with_suffix(".docx")
        if proc.returncode != 0 or not converted.exists():
            result.parse_status = "legacy_doc_unparsed"
            result.notes.append("LibreOffice 转换失败，未产出 docx")
            return
        result.text = _extract_docx_text(converted.read_bytes())
        result.notes.append("经 LibreOffice 转换后按 DOCX 解析")


def _fmt_cell(value: object) -> str:
    """单元格显示值：日期 ISO 化，其余 str()；None（含无缓存值的公式）记空。"""
    if value is None:
        return ""
    if isinstance(value, datetime):
        return (
            value.date().isoformat() if value.time() == datetime.min.time() else value.isoformat()
        )
    if isinstance(value, date):
        return value.isoformat()
    return str(value).strip()


def _sheet_to_lines(ws: object, fill_merged: bool) -> tuple[list[str], bool]:
    """把单个工作表转成 ' | ' 拼接行：跳全空行、合并区域填左上角值、连续重复行去重。

    返回 (行列表, 是否被截断)。
    """
    # 合并单元格填充表：(row, col) -> 区域左上角值（普通模式才有 merged_cells）
    fill: dict[tuple[int, int], object] = {}
    if fill_merged:
        for rng in ws.merged_cells.ranges:
            anchor = ws.cell(row=rng.min_row, column=rng.min_col).value
            for r in range(rng.min_row, rng.max_row + 1):
                for c in range(rng.min_col, rng.max_col + 1):
                    if (r, c) != (rng.min_row, rng.min_col):
                        fill[(r, c)] = anchor
    lines, prev, truncated = [], None, False
    for row in ws.iter_rows():
        cells = [_fmt_cell(fill.get((c.row, c.column), c.value)) for c in row]
        # 去掉行尾空单元格再判全空行
        while cells and not cells[-1]:
            cells.pop()
        if not cells:
            continue
        line = " | ".join(cells)
        if line == prev:  # 连续重复表头行去重
            continue
        if len(lines) >= _MAX_SHEET_ROWS:
            truncated = True
            break
        lines.append(line)
        prev = line
    return lines, truncated


def _extract_xlsx(result: AttachmentParseResult, content: bytes) -> None:
    """多工作表逐个转 Markdown 表格（'## 工作表名' 标题 + ' | ' 拼接行）。"""
    import openpyxl

    # 普通模式才能拿到 merged_cells 做合并填充；大文件退化为 read_only 兜底
    read_only = len(content) > _READ_ONLY_BYTES
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=read_only, data_only=True)
    try:
        result.sheet_count = len(wb.sheetnames)
        if read_only:
            result.notes.append("大文件以 read_only 模式解析，合并单元格未填充")
        parts = []
        for name in wb.sheetnames:
            lines, truncated = _sheet_to_lines(wb[name], fill_merged=not read_only)
            if not lines:
                continue
            parts.append("\n".join([f"## {name}", *lines]))
            if truncated:
                result.notes.append(f"工作表 {name} 超过 {_MAX_SHEET_ROWS} 行，已截断")
        result.text = "\n\n".join(parts).strip()
    finally:
        wb.close()


def parse_attachment(filename: str, content: bytes) -> AttachmentParseResult:
    """解析附件字节流。任何异常都不抛出，统一收敛为 parse_failed + notes 记录原因。"""
    result = AttachmentParseResult(
        filename=filename,
        fmt="unknown",
        text="",
        char_count=0,
        page_count=None,
        sheet_count=None,
        parse_status="ok",
        notes=[],
    )
    try:
        result.fmt = _detect_format(filename, content)
        if result.fmt == "pdf":
            _extract_pdf(result, content)
        elif result.fmt == "docx":
            result.text = _extract_docx_text(content)
        elif result.fmt == "doc":
            _extract_doc(result, content)
        elif result.fmt == "xlsx":
            _extract_xlsx(result, content)
        else:
            # xls/ppt/pptx/unknown：识别格式但不解析内容
            result.parse_status = "unsupported_format"
            result.notes.append(f"暂不支持解析 {result.fmt} 格式的内容")
    except Exception as e:
        result.parse_status = "parse_failed"
        result.text = ""
        result.notes.append(f"解析异常: {type(e).__name__}: {e}")
    result.char_count = len(result.text)
    return result
